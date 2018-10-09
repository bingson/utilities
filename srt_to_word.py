#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import docx
import re
import codecs
from unzip_files import unzip_all 
 
### ToDo
# Add support for multiple directories of .srt files
# Concatenate .docx file generated by each directory of .srt files into a single .docx file
# Add support for a .docx template (a .docx (not .dotx) file containing all the styles you intend to use) if one is present
# Error handling if: no directory with .srt files found, .docx being saved to is in use
# Switch to with open as file pattern if worthwhile
# Multi-platform file reading support
# Sturdier way to skip lines containing srt line number
# switch to logging over print
# text files for a single directory contain


#rejection pattern matches
#valid_lines = re.compile(r'[a-z|A-Z|/.$]') # has letters in the line or ends in a period
valid_lines = re.compile(r'[a-z|A-Z]|/.$') # has letters in the line or ends in a period

def clean_files(file_directory):
    
    os.chdir(file_directory) # changes directory  
    
    
    file_list = os.listdir(file_directory) # gets  list of files in \\srt folder
    doc = docx.Document()  # creates document object that will receive all the srt lines
    doc_master = docx.Document()
    
    
#    toc_link = "###### [CONTENTS](#CONTENTS)\n" # table of contents
    toc_link = ""
    
    count = 0
    
    for file_name in file_list:
        try:
            if file_name.endswith('.srt'):
                count += 1
                input_file = codecs.open(file_name, encoding='utf-8')  # https://docs.python.org/2.7/howto/unicode.html
                work_space = input_file.readlines()
            else: continue
    
            # Create word document objects
            toc_entry = "\n[" + file_name[:-4] +"]" + "(#"+ file_name[:2] +")<br>\n"
            toc_anchor = "<a id='" + file_name[:2] + "'></a>" 
            
            file_header = '\n'+'## '+ str(file_name)[:-4] + str(toc_anchor) + toc_entry + toc_link
            # str(file_name)[:-4] removes file extension to create h2 title
            
#           print file_header
            doc.add_heading(file_header, 5) # adds file_header with word heading level 5
            txt_paragraph = ''              # creates an empty string that will receive each line of srt data
#            paragraph = doc
            paragraph = doc.add_paragraph() # add blank paragraph to doc to separate different srt files
    
            for line in work_space:
                if not re.search(valid_lines, line): continue # skips lines we don't want
                else: txt_paragraph += line.replace('\r', '').replace('\n', '').replace('>>', '') + ' ' 
                    # .strip('>>')
                    # what makes it into doc
                    # creates paragraph that will later be added to doc
                    # take all new line characters out \r and \n (two types rlated to windows), 
                    # then adds a space at the end
            paragraph.add_run(txt_paragraph.replace('  ',' '))
            input_file.close() # closes srt file being used so it can open another one
            
        except: print file_name
    doc_master.add_heading(toc_entry, 5)
    doc.save('Cleaned.docx')

file_directory = "C:\Users\\Bing\\Google Drive\\1-Data Science\\srt"

unzip_all(file_directory)
clean_files(file_directory)